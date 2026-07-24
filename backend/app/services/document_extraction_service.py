"""Document text extraction, classification, and structured field parsing.

Supports profile auto-fill documents only:
  - National ID (CNIC/NIC)
  - Passport (foreigners)
  - Resume / CV
  - Academic Transcript

Classification is content-based (not filename). Callers pass an expected
purpose/category so wrong document types are rejected before autofill.
"""

from __future__ import annotations

import json
import logging
import os
import re
import tempfile
from pathlib import Path

from app.core.config import settings
from app.schemas.document import PURPOSE_EXPECTED_CATEGORIES, PURPOSE_REJECT_MESSAGES

logger = logging.getLogger("document_extraction_service")

CNIC_PATTERN = re.compile(r"\b\d{5}-?\d{7}-?\d{1}\b")
PASSPORT_PATTERN = re.compile(r"\b[A-Z]{1,2}\d{6,8}\b")
EMAIL_PATTERN = re.compile(r"[\w.+-]+@[\w.-]+\.\w+")
PHONE_PATTERN = re.compile(r"\b(?:\+?92|0)?3\d{2}[- ]?\d{7}\b")
DATE_PATTERN = re.compile(r"\b(\d{1,2}[-/ ]\d{1,2}[-/ ]\d{2,4}|\d{4}[-/ ]\d{1,2}[-/ ]\d{1,2})\b")
GPA_PATTERN = re.compile(r"\b(?:cgpa|gpa|g\.?p\.?a\.?)\s*[:=]?\s*(\d+(?:\.\d+)?)\b", re.I)
PERCENT_PATTERN = re.compile(r"\b(\d{2,3}(?:\.\d+)?)\s*%")
YEAR_PATTERN = re.compile(r"\b((?:19|20)\d{2})\b")

DRIVING_LICENSE_HINTS = ("driving licence", "driving license", "dl no", "learner permit")
TRANSCRIPT_HINTS = ("transcript", "mark sheet", "marks sheet", "grade sheet", "academic record", "semester")
RESUME_HINTS = ("curriculum vitae", " curriculum", "experience", "work history", "skills", "objective", "summary")
CNIC_HINTS = ("nadra", "cnic", "national identity", "identity card", "father name", "gender")
PASSPORT_HINTS = ("passport", "nationality", "place of birth", "surname", "given names")

# Header / label noise that must never become a person's name.
CNIC_NAME_REJECT = frozenset(
    {
        "pakistan",
        "islamic republic of pakistan",
        "national identity card",
        "identity card",
        "govt of pakistan",
        "government of pakistan",
        "country of stay",
        "holder's signature",
        "holders signature",
        "holder signature",
        "name",
        "father name",
        "father's name",
        "gender",
        "identity number",
        "date of birth",
        "date of issue",
        "date of expiry",
        "nationality",
        "nic",
        "cnic",
    }
)

PK_CNIC_DATE_PATTERN = re.compile(r"\b(\d{1,2}[.\-/]\d{1,2}[.\-/]\d{4})\b")

# --- Banking (cheque / bank letter / account maintenance certificate) --------
IBAN_TEXT_PATTERN = re.compile(r"\bPK\s?\d{2}(?:[ -]?[A-Z0-9]){16,20}\b", re.I)
ACCOUNT_NUMBER_PATTERN = re.compile(r"\b\d[\d-]{7,23}\d\b")
SWIFT_PATTERN = re.compile(r"\b[A-Z]{4}PK[A-Z0-9]{2}(?:[A-Z0-9]{3})?\b")
BRANCH_CODE_PATTERN = re.compile(r"\b(?:branch\s*(?:code|#|no\.?)\s*[:\-]?\s*)(\d{2,6})\b", re.I)
NTN_PATTERN = re.compile(r"\b(?:ntn|tax\s*(?:id|number))\s*[:\-]?\s*([\d-]{6,15})\b", re.I)

BANK_DOCUMENT_HINTS = (
    "iban",
    "account title",
    "account holder",
    "account number",
    "branch code",
    "swift",
    "cheque",
    "bank statement",
    "account maintenance",
)

# Well-known Pakistani banks, longest-first so "MCB Islamic" wins over "MCB".
PK_BANK_NAMES = (
    "Allied Bank Limited",
    "Askari Bank",
    "Bank Alfalah",
    "Bank Al Habib",
    "BankIslami",
    "Dubai Islamic Bank",
    "Faysal Bank",
    "Habib Bank Limited",
    "Habib Metropolitan Bank",
    "JS Bank",
    "MCB Islamic Bank",
    "MCB Bank",
    "Meezan Bank",
    "National Bank of Pakistan",
    "Samba Bank",
    "Silkbank",
    "Sindh Bank",
    "Soneri Bank",
    "Standard Chartered",
    "Summit Bank",
    "The Bank of Punjab",
    "United Bank Limited",
    "Allied Bank",
    "HBL",
    "UBL",
    "ABL",
    "NBP",
)

BANK_FIELD_KEYS = (
    "bank_name",
    "account_holder_name",
    "account_number",
    "iban",
    "branch",
    "branch_code",
    "swift_code",
    "tax_id",
)


class DocumentExtractionService:
    def __init__(self):
        self._easyocr_reader = None

    def _get_easyocr_reader(self):
        if self._easyocr_reader is None:
            try:
                import easyocr

                self._easyocr_reader = easyocr.Reader(["en"], gpu=False)
            except Exception as e:
                logger.error(f"Failed to initialize EasyOCR: {e}")
        return self._easyocr_reader

    def detect_file_type(self, file_path: str) -> str:
        ext = Path(file_path).suffix.lower()
        if ext in (".png", ".jpg", ".jpeg", ".webp", ".bmp", ".tif", ".tiff"):
            return "image"
        if ext == ".pdf":
            return "pdf"
        if ext in (".doc", ".docx"):
            return "docx"
        if ext in (".txt", ".log", ".json", ".csv"):
            return "txt"
        # Cloudinary raw URLs sometimes strip extensions — sniff magic bytes.
        try:
            with open(file_path, "rb") as handle:
                header = handle.read(8)
            if header.startswith(b"%PDF"):
                return "pdf"
            if header.startswith(b"\x89PNG") or header[:2] == b"\xff\xd8":
                return "image"
            if header[:2] == b"PK":
                return "docx"
        except OSError:
            pass
        return "unknown"

    def extract_text_from_txt(self, file_path: str) -> str:
        try:
            with open(file_path, "r", encoding="utf-8") as f:
                return f.read()
        except UnicodeDecodeError:
            with open(file_path, "r", encoding="latin-1") as f:
                return f.read()

    def extract_text_from_docx(self, file_path: str) -> str:
        import docx

        doc = docx.Document(file_path)
        parts = []
        for p in doc.paragraphs:
            if p.text.strip():
                parts.append(p.text)
        for table in doc.tables:
            for row in table.rows:
                cells = [c.text.strip() for c in row.cells if c.text.strip()]
                if cells:
                    parts.append(" | ".join(cells))
        return "\n".join(parts)

    def extract_text_from_image(self, file_path: str) -> str:
        reader = self._get_easyocr_reader()
        extracted_lines = []
        if reader:
            try:
                results = reader.readtext(file_path, detail=0)
                if results:
                    extracted_lines = results
            except Exception as e:
                logger.error(f"EasyOCR extraction failed: {e}")

        text = "\n".join(extracted_lines)
        if text.strip():
            return text

        logger.info("Falling back to pytesseract for image OCR")
        try:
            import pytesseract
            from PIL import Image

            tess_text = pytesseract.image_to_string(Image.open(file_path))
            if tess_text.strip():
                return tess_text
        except Exception as e:
            logger.error(f"Pytesseract fallback failed: {e}")

        return text

    def extract_text_from_pdf(self, file_path: str) -> str:
        import fitz

        doc = fitz.open(file_path)
        text_parts = []
        is_scanned = True

        for page in doc:
            page_text = page.get_text()
            if page_text.strip():
                text_parts.append(page_text)
                is_scanned = False

        if not is_scanned:
            doc.close()
            return "\n".join(text_parts)

        logger.info(f"PDF {file_path} appears to be scanned. Rasterizing pages...")
        reader = self._get_easyocr_reader()
        for page_num in range(len(doc)):
            page = doc.load_page(page_num)
            pix = page.get_pixmap(dpi=150)

            with tempfile.NamedTemporaryFile(suffix=".png", delete=False) as tmp:
                tmp_name = tmp.name

            try:
                pix.save(tmp_name)
                page_text = ""
                if reader:
                    try:
                        results = reader.readtext(tmp_name, detail=0)
                        page_text = " ".join(results)
                    except Exception as e:
                        logger.error(f"EasyOCR on PDF page {page_num} failed: {e}")

                if not page_text.strip():
                    try:
                        import pytesseract
                        from PIL import Image

                        page_text = pytesseract.image_to_string(Image.open(tmp_name))
                    except Exception as e:
                        logger.error(f"Pytesseract fallback on PDF page {page_num} failed: {e}")

                if page_text.strip():
                    text_parts.append(page_text)
            finally:
                if os.path.exists(tmp_name):
                    try:
                        os.unlink(tmp_name)
                    except OSError:
                        pass

        doc.close()
        return "\n".join(text_parts)

    def extract_text(self, file_path: str) -> str:
        file_type = self.detect_file_type(file_path)
        if file_type == "pdf":
            return self.extract_text_from_pdf(file_path)
        if file_type == "docx":
            return self.extract_text_from_docx(file_path)
        if file_type == "txt":
            return self.extract_text_from_txt(file_path)
        if file_type == "image":
            return self.extract_text_from_image(file_path)
        return ""

    @staticmethod
    def resolve_expected_categories(
        *,
        purpose: str | None = None,
        doc_type: str | None = None,
        category: str | None = None,
    ) -> tuple[str, ...]:
        """Map upload purpose/doc_type to accepted classification categories."""
        if doc_type in PURPOSE_EXPECTED_CATEGORIES:
            return PURPOSE_EXPECTED_CATEGORIES[doc_type]
        if purpose in PURPOSE_EXPECTED_CATEGORIES:
            expected = PURPOSE_EXPECTED_CATEGORIES[purpose]
            # Narrow government_doc when the user picked CNIC or Passport explicitly.
            if purpose == "government_doc" and doc_type in ("cnic", "passport"):
                return (doc_type,)
            return expected
        if category == "identity":
            if doc_type in ("cnic", "passport"):
                return (doc_type,)
            return PURPOSE_EXPECTED_CATEGORIES["identity"]
        if category == "education" or doc_type in ("transcript", "degree"):
            return PURPOSE_EXPECTED_CATEGORIES["transcript"]
        if doc_type == "resume" or category == "other":
            if doc_type == "resume":
                return PURPOSE_EXPECTED_CATEGORIES["resume"]
        return ()

    @staticmethod
    def reject_message_for(
        *,
        purpose: str | None = None,
        doc_type: str | None = None,
        expected: tuple[str, ...] = (),
    ) -> str:
        if doc_type in PURPOSE_REJECT_MESSAGES:
            return PURPOSE_REJECT_MESSAGES[doc_type]
        if purpose in PURPOSE_REJECT_MESSAGES:
            if purpose == "government_doc" and expected == ("cnic",):
                return PURPOSE_REJECT_MESSAGES["cnic"]
            if purpose == "government_doc" and expected == ("passport",):
                return PURPOSE_REJECT_MESSAGES["passport"]
            return PURPOSE_REJECT_MESSAGES[purpose]
        if expected == ("cnic",):
            return PURPOSE_REJECT_MESSAGES["cnic"]
        if expected == ("passport",):
            return PURPOSE_REJECT_MESSAGES["passport"]
        if expected == ("resume",):
            return PURPOSE_REJECT_MESSAGES["resume"]
        if "academic_transcript" in expected:
            return PURPOSE_REJECT_MESSAGES["transcript"]
        return "Uploaded document does not match the required document type."

    def validate_classification(
        self,
        parsed: dict,
        expected: tuple[str, ...],
        *,
        purpose: str | None = None,
        doc_type: str | None = None,
    ) -> dict:
        """Return validation result. Rejects wrong types (passport as CNIC, resume as ID, etc.)."""
        category = (parsed or {}).get("category") or "unknown"
        classification_confidence = float((parsed or {}).get("classification_confidence") or 0.0)

        if not expected:
            return {
                "accepted": True,
                "category": category,
                "classification_confidence": classification_confidence,
                "rejection_message": None,
            }

        # Driving license / unknown never accepted for profile autofill slots.
        if category in ("driving_license", "unknown", "payroll", "invoice"):
            return {
                "accepted": False,
                "category": category,
                "classification_confidence": classification_confidence,
                "rejection_message": self.reject_message_for(
                    purpose=purpose, doc_type=doc_type, expected=expected
                ),
            }

        if category not in expected:
            return {
                "accepted": False,
                "category": category,
                "classification_confidence": classification_confidence,
                "rejection_message": self.reject_message_for(
                    purpose=purpose, doc_type=doc_type, expected=expected
                ),
            }

        return {
            "accepted": True,
            "category": category,
            "classification_confidence": classification_confidence,
            "rejection_message": None,
        }

    @staticmethod
    def _field_completeness(fields: dict) -> float:
        if not fields:
            return 0.0
        values = list(fields.values())
        if not values:
            return 0.0
        filled = 0
        for v in values:
            if v is None or v == "" or v == [] or v == {}:
                continue
            filled += 1
        return round(filled / len(values), 4)

    async def parse_structured_data(self, raw_text: str) -> dict:
        """Parse structured fields using Gemini API, degrading to heuristics."""
        if not raw_text.strip():
            return {
                "category": "unknown",
                "fields": {},
                "classification_confidence": 0.0,
                "extraction_confidence": 0.0,
            }

        if settings.OPENROUTER_API_KEY or (
            settings.GEMINI_API_KEY and not settings.GEMINI_API_KEY.strip().startswith("YOUR_")
        ):
            try:
                from app.services.llm_service import call_llm_json

                prompt = f"""
You are an expert document classifier and information extractor.
Analyze the document text below.

First classify into ONE category:
- "cnic" — Pakistan National Identity Card / NIC / National ID only
- "passport" — Passport biodata page
- "resume" — Resume or CV
- "academic_transcript" — Academic transcript / mark sheet / grade sheet
- "certificate" — Degree certificate (not a transcript)
- "driving_license" — Driving licence
- "payroll" — Salary slip / payroll
- "invoice" — Invoice / bill
- "unknown" — Anything else

Also return classification_confidence (0.0–1.0).

Then extract fields for that category only. Missing fields = null.

1. "cnic":
   name, first_name, last_name, father_name, cnic_number, date_of_birth,
   gender, nationality, marital_status, issue_date, expiry_date

2. "passport":
   name, first_name, last_name, passport_number, nationality, date_of_birth,
   gender, place_of_birth, issue_date, expiry_date

3. "resume":
   full_name, first_name, last_name, email, phone_number, address,
   professional_summary, technical_skills (list), soft_skills (list),
   languages (list), skills (list — all skills if not split),
   work_experience (list of {{job_title, company, start_date, end_date, description}}),
   projects (list of strings or objects), achievements (list),
   education (list of {{institute, degree, program, major, year}}),
   certifications (list), awards (list),
   linkedin, github, portfolio

4. "academic_transcript":
   candidate_name, institute, degree, program, major,
   cgpa, gpa, percentage, passing_year, subjects (list of strings or objects)

5. "certificate" (degree / education certificate — still extract academic fields):
   candidate_name, institute, degree, program, major,
   cgpa, gpa, percentage, passing_year, issue_date

6. Other categories: extract whatever identity-like fields are obvious.

Return JSON only:
{{
  "category": "...",
  "classification_confidence": 0.0,
  "fields": {{ ... }}
}}

Document Text:
{raw_text[:12000]}
"""

                parsed = await call_llm_json(prompt, timeout=45.0)
                if parsed and "category" in parsed and "fields" in parsed:
                    fields = parsed.get("fields") or {}
                    if parsed.get("category") == "cnic":
                        fields = self._sanitize_cnic_fields(fields, raw_text)
                    class_conf = float(parsed.get("classification_confidence") or 0.85)
                    return {
                        "category": parsed["category"],
                        "fields": fields,
                        "classification_confidence": class_conf,
                        "extraction_confidence": self._field_completeness(fields),
                    }

            except Exception as e:
                logger.error(f"LLM document parse failed: {e}. Falling back to heuristics.")

        return self._heuristic_fallback_parse(raw_text)

    def _heuristic_fallback_parse(self, text: str) -> dict:
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        lower = text.lower()

        cnic_match = CNIC_PATTERN.search(text)
        passport_match = PASSPORT_PATTERN.search(text)
        email_match = EMAIL_PATTERN.search(text)
        phone_match = PHONE_PATTERN.search(text)
        dates = DATE_PATTERN.findall(text)
        gpa_match = GPA_PATTERN.search(text)
        percent_match = PERCENT_PATTERN.search(text)
        years = YEAR_PATTERN.findall(text)

        def has_any(hints: tuple[str, ...]) -> bool:
            return any(h in lower for h in hints)

        # Prefer explicit negative classes first.
        if has_any(DRIVING_LICENSE_HINTS):
            return {
                "category": "driving_license",
                "fields": {},
                "classification_confidence": 0.8,
                "extraction_confidence": 0.0,
            }

        if cnic_match or (has_any(CNIC_HINTS) and not has_any(PASSPORT_HINTS)):
            fields = self._sanitize_cnic_fields(self._parse_pk_cnic_fields(text), text)
            return {
                "category": "cnic",
                "fields": fields,
                "classification_confidence": 0.75 if cnic_match else 0.55,
                "extraction_confidence": self._field_completeness(fields),
            }

        if passport_match or has_any(PASSPORT_HINTS):
            name = lines[0] if lines else None
            first_name, last_name = self._split_name(name)
            fields = {
                "name": name,
                "first_name": first_name,
                "last_name": last_name,
                "passport_number": passport_match.group(0) if passport_match else None,
                "nationality": None,
                "date_of_birth": dates[0] if dates else None,
                "gender": "male" if "male" in lower else "female" if "female" in lower else None,
                "place_of_birth": None,
                "issue_date": dates[1] if len(dates) > 1 else None,
                "expiry_date": dates[-1] if dates else None,
            }
            return {
                "category": "passport",
                "fields": fields,
                "classification_confidence": 0.7 if passport_match else 0.5,
                "extraction_confidence": self._field_completeness(fields),
            }

        if has_any(TRANSCRIPT_HINTS) or (gpa_match and has_any(("university", "college", "semester", "cgpa"))):
            fields = {
                "candidate_name": lines[0] if lines else None,
                "institute": next(
                    (l for l in lines if any(k in l.lower() for k in ("university", "college", "institute", "school"))),
                    None,
                ),
                "degree": next(
                    (l for l in lines if any(k in l.lower() for k in ("bachelor", "master", "bs", "ms", "phd", "degree"))),
                    None,
                ),
                "program": None,
                "major": None,
                "cgpa": gpa_match.group(1) if gpa_match else None,
                "gpa": gpa_match.group(1) if gpa_match else None,
                "percentage": percent_match.group(1) if percent_match else None,
                "passing_year": years[-1] if years else None,
                "subjects": [],
            }
            return {
                "category": "academic_transcript",
                "fields": fields,
                "classification_confidence": 0.7,
                "extraction_confidence": self._field_completeness(fields),
            }

        if email_match or has_any(RESUME_HINTS):
            name = lines[0] if lines else None
            first_name, last_name = self._split_name(name)
            fields = {
                "full_name": name,
                "first_name": first_name,
                "last_name": last_name,
                "email": email_match.group(0) if email_match else None,
                "phone_number": phone_match.group(0) if phone_match else None,
                "address": None,
                "professional_summary": None,
                "technical_skills": [],
                "soft_skills": [],
                "languages": [],
                "skills": [],
                "work_experience": [],
                "projects": [],
                "achievements": [],
                "education": [],
                "certifications": [],
                "awards": [],
                "linkedin": None,
                "github": None,
                "portfolio": None,
            }
            return {
                "category": "resume",
                "fields": fields,
                "classification_confidence": 0.65 if email_match else 0.5,
                "extraction_confidence": self._field_completeness(fields),
            }

        if "salary" in lower or "pay slip" in lower or "payroll" in lower:
            return {
                "category": "payroll",
                "fields": {},
                "classification_confidence": 0.6,
                "extraction_confidence": 0.0,
            }

        if "invoice" in lower or "tax invoice" in lower:
            return {
                "category": "invoice",
                "fields": {},
                "classification_confidence": 0.6,
                "extraction_confidence": 0.0,
            }

        if "degree" in lower or "certificate" in lower:
            fields = {
                "candidate_name": lines[0] if lines else None,
                "institute": next(
                    (l for l in lines if any(k in l.lower() for k in ("university", "college", "institute"))),
                    None,
                ),
                "degree": next(
                    (l for l in lines if any(k in l.lower() for k in ("bachelor", "master", "degree"))),
                    None,
                ),
                "completion_date": dates[-1] if dates else None,
            }
            return {
                "category": "certificate",
                "fields": fields,
                "classification_confidence": 0.55,
                "extraction_confidence": self._field_completeness(fields),
            }

        return {
            "category": "unknown",
            "fields": {},
            "classification_confidence": 0.2,
            "extraction_confidence": 0.0,
        }

    @staticmethod
    def _split_name(full_name: str | None) -> tuple[str | None, str | None]:
        if not full_name:
            return None, None
        parts = [p for p in str(full_name).strip().split() if p]
        if not parts:
            return None, None
        if len(parts) == 1:
            return parts[0], None
        return parts[0], " ".join(parts[1:])

    @staticmethod
    def _is_invalid_person_name(value: str | None) -> bool:
        if not value:
            return True
        cleaned = re.sub(r"\s+", " ", str(value).strip())
        if len(cleaned) < 2:
            return True
        lower = cleaned.lower()
        if lower in CNIC_NAME_REJECT:
            return True
        if lower.startswith("islamic republic") or lower.startswith("national identity"):
            return True
        if re.fullmatch(r"[\d\W_]+", cleaned):
            return True
        if cleaned.isupper() and cleaned in {"PAKISTAN", "NADRA"}:
            return True
        return False

    @staticmethod
    def _normalize_gender_token(value: str | None) -> str | None:
        if not value:
            return None
        token = str(value).strip().lower()
        if token in {"m", "male"}:
            return "male"
        if token in {"f", "female"}:
            return "female"
        return None

    @staticmethod
    def _parse_pk_date_sort_key(raw: str) -> tuple[int, int, int] | None:
        match = re.match(r"^(\d{1,2})[.\-/](\d{1,2})[.\-/](\d{4})$", str(raw).strip())
        if not match:
            return None
        day, month, year = (int(match.group(1)), int(match.group(2)), int(match.group(3)))
        if not (1 <= month <= 12 and 1 <= day <= 31 and 1900 <= year <= 2100):
            return None
        return year, month, day

    def _value_after_label(self, lines: list[str], label: str, *, stop_labels: tuple[str, ...] = ()) -> str | None:
        """Read CNIC field value on the same line as a label or on the next line."""
        label_lower = label.lower()
        stop = {s.lower() for s in stop_labels}

        for index, line in enumerate(lines):
            stripped = line.strip()
            lower = stripped.lower()

            if label_lower not in lower:
                continue

            # Same-line value: "Name: Faris Altaf Dosani" or "Name Faris Altaf Dosani"
            same_line = re.sub(rf"^{re.escape(label)}\s*[:\-]?\s*", "", stripped, flags=re.I).strip()
            if same_line and same_line.lower() != label_lower and same_line.lower() not in stop:
                return same_line

            # Next-line value (common on NADRA cards)
            for next_index in range(index + 1, min(index + 3, len(lines))):
                candidate = lines[next_index].strip()
                candidate_lower = candidate.lower()
                if not candidate:
                    continue
                if candidate_lower in stop or any(stop_label in candidate_lower for stop_label in stop):
                    break
                if candidate_lower in CNIC_NAME_REJECT and label_lower == "name":
                    continue
                return candidate
        return None

    def _parse_pk_cnic_fields(self, text: str) -> dict:
        """Label-aware parser for Pakistan CNIC layout (NADRA cards)."""
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        lower_text = text.lower()

        cnic_match = CNIC_PATTERN.search(text)
        cnic_number = cnic_match.group(0) if cnic_match else None

        name = self._value_after_label(
            lines,
            "Name",
            stop_labels=("Father Name", "Father's Name", "Gender", "Identity Number", "Date of Birth"),
        )
        father_name = self._value_after_label(
            lines,
            "Father Name",
            stop_labels=("Gender", "Identity Number", "Date of Birth", "Date of Issue", "Date of Expiry"),
        )
        if not father_name:
            father_name = self._value_after_label(
                lines,
                "Father's Name",
                stop_labels=("Gender", "Identity Number", "Date of Birth", "Date of Issue", "Date of Expiry"),
            )

        gender_raw = self._value_after_label(lines, "Gender", stop_labels=("Country of Stay", "Identity Number"))
        gender = self._normalize_gender_token(gender_raw)
        if not gender:
            gender_match = re.search(r"\bGender\s*[:\-]?\s*([MF])\b", text, re.I)
            if gender_match:
                gender = self._normalize_gender_token(gender_match.group(1))

        # Labelled dates first
        date_of_birth = self._value_after_label(lines, "Date of Birth", stop_labels=("Date of Issue",))
        issue_date = self._value_after_label(lines, "Date of Issue", stop_labels=("Date of Expiry",))
        expiry_date = self._value_after_label(lines, "Date of Expiry")

        # Fallback: assign chronological PK-style dates when labels are missing
        if not (date_of_birth and issue_date and expiry_date):
            dated = []
            for match in PK_CNIC_DATE_PATTERN.finditer(text):
                raw = match.group(1)
                sort_key = self._parse_pk_date_sort_key(raw)
                if sort_key and raw not in {d[1] for d in dated}:
                    dated.append((sort_key, raw))
            dated.sort(key=lambda item: item[0])
            if not date_of_birth and len(dated) >= 1:
                date_of_birth = dated[0][1]
            if not issue_date and len(dated) >= 2:
                issue_date = dated[1][1]
            if not expiry_date and len(dated) >= 3:
                expiry_date = dated[2][1]

        if self._is_invalid_person_name(name):
            # Try line immediately after a standalone "Name" label block in noisy OCR.
            for index, line in enumerate(lines):
                if line.strip().lower() == "name" and index + 1 < len(lines):
                    candidate = lines[index + 1].strip()
                    if not self._is_invalid_person_name(candidate):
                        name = candidate
                        break

        first_name, last_name = self._split_name(name)
        nationality = "Pakistani" if "pakistan" in lower_text else None

        fields = {
            "name": name if not self._is_invalid_person_name(name) else None,
            "first_name": first_name if not self._is_invalid_person_name(first_name) else None,
            "last_name": last_name,
            "father_name": father_name if father_name and father_name.lower() not in CNIC_NAME_REJECT else None,
            "cnic_number": cnic_number,
            "date_of_birth": date_of_birth,
            "gender": gender,
            "nationality": nationality,
            "marital_status": None,
            "issue_date": issue_date,
            "expiry_date": expiry_date,
        }
        return fields

    def _sanitize_cnic_fields(self, fields: dict, raw_text: str) -> dict:
        """Fix common CNIC extraction mistakes and backfill from label parser."""
        cleaned = dict(fields or {})
        parsed = self._parse_pk_cnic_fields(raw_text)

        for key in (
            "name",
            "first_name",
            "last_name",
            "father_name",
            "cnic_number",
            "date_of_birth",
            "gender",
            "nationality",
            "issue_date",
            "expiry_date",
        ):
            current = cleaned.get(key)
            fallback = parsed.get(key)
            if key in {"name", "first_name"} and self._is_invalid_person_name(current):
                cleaned[key] = fallback
            elif not current and fallback:
                cleaned[key] = fallback

        if self._is_invalid_person_name(cleaned.get("name")) and cleaned.get("first_name"):
            cleaned["name"] = " ".join(
                part for part in (cleaned.get("first_name"), cleaned.get("last_name")) if part
            ).strip() or None

        if cleaned.get("name") and not cleaned.get("first_name"):
            first_name, last_name = self._split_name(cleaned.get("name"))
            cleaned["first_name"] = first_name
            cleaned["last_name"] = last_name

        cleaned["gender"] = self._normalize_gender_token(cleaned.get("gender")) or parsed.get("gender")
        return cleaned

    # ------------------------------------------------------------------
    # Banking documents (cheque, bank letter, account maintenance cert)
    # ------------------------------------------------------------------
    @staticmethod
    def _normalize_iban(value: str | None) -> str | None:
        if not value:
            return None
        cleaned = re.sub(r"[^A-Z0-9]", "", str(value).upper())
        return cleaned or None

    def parse_bank_document_sync(self, raw_text: str) -> dict:
        """Regex/keyword extraction of banking fields. Always available offline."""
        text = raw_text or ""
        lines = [ln.strip() for ln in text.splitlines() if ln.strip()]
        lower = text.lower()

        fields: dict = {key: None for key in BANK_FIELD_KEYS}
        alternatives: dict = {}

        iban_matches = []
        for match in IBAN_TEXT_PATTERN.finditer(text):
            normalized = self._normalize_iban(match.group(0))
            if normalized and 20 <= len(normalized) <= 24 and normalized not in iban_matches:
                iban_matches.append(normalized)
        if iban_matches:
            fields["iban"] = iban_matches[0]
            if len(iban_matches) > 1:
                alternatives["iban"] = iban_matches[1:4]

        bank_hits = [name for name in PK_BANK_NAMES if name.lower() in lower]
        if bank_hits:
            fields["bank_name"] = bank_hits[0]
            if len(bank_hits) > 1:
                alternatives["bank_name"] = bank_hits[1:4]

        title_candidates = []
        for label in ("account title", "account holder", "title of account", "customer name", "a/c title"):
            value = self._value_after_label(lines, label)
            if value and value not in title_candidates:
                title_candidates.append(value)
        if title_candidates:
            fields["account_holder_name"] = title_candidates[0]
            if len(title_candidates) > 1:
                alternatives["account_holder_name"] = title_candidates[1:4]

        account_candidates = []
        for label in ("account number", "account no", "a/c no", "a/c #", "account #"):
            value = self._value_after_label(lines, label)
            digits = re.sub(r"[^\d-]", "", value or "")
            if digits and len(re.sub(r"\D", "", digits)) >= 8 and digits not in account_candidates:
                account_candidates.append(digits)
        if not account_candidates and fields["iban"]:
            # Pakistani IBANs end with the zero-padded account number.
            tail = fields["iban"][8:].lstrip("0")
            if len(tail) >= 8:
                account_candidates.append(tail)
        if account_candidates:
            fields["account_number"] = account_candidates[0]
            if len(account_candidates) > 1:
                alternatives["account_number"] = account_candidates[1:4]

        branch = self._value_after_label(lines, "branch name") or self._value_after_label(lines, "branch")
        if branch and not branch.isdigit() and not branch.lower().startswith("code"):
            fields["branch"] = branch

        branch_code_match = BRANCH_CODE_PATTERN.search(text)
        if branch_code_match:
            fields["branch_code"] = branch_code_match.group(1)

        swift_match = SWIFT_PATTERN.search(text.upper())
        if swift_match:
            fields["swift_code"] = swift_match.group(0)

        ntn_match = NTN_PATTERN.search(text)
        if ntn_match:
            fields["tax_id"] = ntn_match.group(1)

        is_bank_document = bool(
            fields["iban"] or fields["bank_name"] or any(hint in lower for hint in BANK_DOCUMENT_HINTS)
        )

        return {
            "fields": fields,
            "alternatives": alternatives,
            "is_bank_document": is_bank_document,
            "engine": "heuristic",
        }

    async def parse_bank_document(self, raw_text: str) -> dict:
        """Extract banking fields from a cheque / bank letter, with confidences.

        The LLM handles layout variety; the regex pass runs either way and is
        used to backfill gaps and to validate the IBAN the model returned.
        """
        heuristic = self.parse_bank_document_sync(raw_text)

        if not (raw_text or "").strip():
            return {
                "fields": {key: None for key in BANK_FIELD_KEYS},
                "field_confidence": {},
                "alternatives": {},
                "extraction_confidence": 0.0,
                "is_bank_document": False,
                "engine": "none",
            }

        llm_fields: dict = {}
        llm_confidence: dict = {}
        llm_alternatives: dict = {}
        engine = heuristic["engine"]
        is_bank_document = heuristic["is_bank_document"]

        if settings.OPENROUTER_API_KEY or (
            settings.GEMINI_API_KEY and not settings.GEMINI_API_KEY.strip().startswith("YOUR_")
        ):
            try:
                from app.services.llm_service import call_llm_json

                prompt = f"""
You are extracting bank account details from a Pakistani banking document
(a cheque, bank letter, or account maintenance certificate).

Extract these fields. Use null when the document does not clearly show a value.
Never guess an IBAN or account number.

- bank_name: the bank's name, e.g. "Meezan Bank" or "Habib Bank Limited"
- account_holder_name: the account title exactly as printed
- account_number: digits only (hyphens allowed)
- iban: 24-character Pakistani IBAN starting with PK, no spaces
- branch: branch name or address
- branch_code: numeric branch code
- swift_code: SWIFT/BIC code
- tax_id: NTN / tax number if printed

For every field also return a confidence between 0.0 and 1.0 reflecting how
clearly it was printed and how sure you are.

If the document shows more than one plausible value for a field (for example
two different names), list the others under "alternatives".

Set "is_bank_document" to false if this is not a banking document at all.

Return JSON only:
{{
  "is_bank_document": true,
  "fields": {{ "bank_name": null, "account_holder_name": null, "account_number": null,
               "iban": null, "branch": null, "branch_code": null, "swift_code": null, "tax_id": null }},
  "confidence": {{ "bank_name": 0.0 }},
  "alternatives": {{ "account_holder_name": [] }}
}}

Document Text:
{(raw_text or "")[:8000]}
"""
                parsed = await call_llm_json(prompt, timeout=45.0)
                if parsed and isinstance(parsed.get("fields"), dict):
                    llm_fields = parsed["fields"]
                    llm_confidence = parsed.get("confidence") or {}
                    llm_alternatives = parsed.get("alternatives") or {}
                    engine = "llm"
                    if parsed.get("is_bank_document") is True:
                        is_bank_document = True
                    elif parsed.get("is_bank_document") is False and not heuristic["fields"]["iban"]:
                        is_bank_document = False
            except Exception as exc:
                logger.error(f"LLM bank document parse failed: {exc}. Falling back to heuristics.")

        fields: dict = {}
        field_confidence: dict = {}
        for key in BANK_FIELD_KEYS:
            llm_value = llm_fields.get(key)
            if isinstance(llm_value, str):
                llm_value = llm_value.strip() or None
            heuristic_value = heuristic["fields"].get(key)

            if key == "iban":
                llm_value = self._normalize_iban(llm_value)
                # Trust the printed IBAN over the model when they disagree.
                if heuristic_value and llm_value and heuristic_value != llm_value:
                    fields[key] = heuristic_value
                    field_confidence[key] = 0.55
                    existing = list(llm_alternatives.get(key) or [])
                    llm_alternatives[key] = [llm_value, *existing]
                    continue

            if llm_value:
                fields[key] = llm_value
                try:
                    field_confidence[key] = max(0.0, min(1.0, float(llm_confidence.get(key, 0.75))))
                except (TypeError, ValueError):
                    field_confidence[key] = 0.75
            elif heuristic_value:
                fields[key] = heuristic_value
                field_confidence[key] = 0.6
            else:
                fields[key] = None

        alternatives = {**heuristic["alternatives"]}
        for key, values in (llm_alternatives or {}).items():
            if key not in BANK_FIELD_KEYS or not isinstance(values, list):
                continue
            merged = [str(v).strip() for v in values if str(v or "").strip()]
            merged.extend(alternatives.get(key, []))
            deduped = []
            for value in merged:
                if value != fields.get(key) and value not in deduped:
                    deduped.append(value)
            if deduped:
                alternatives[key] = deduped[:4]

        # Any field with a listed alternative is genuinely ambiguous.
        for key in alternatives:
            field_confidence[key] = min(field_confidence.get(key, 0.6), 0.55)

        return {
            "fields": fields,
            "field_confidence": field_confidence,
            "alternatives": alternatives,
            "extraction_confidence": self._field_completeness(fields),
            "is_bank_document": is_bank_document,
            "engine": engine,
        }


document_extraction_service = DocumentExtractionService()
